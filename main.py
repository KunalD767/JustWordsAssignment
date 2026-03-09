import os
import requests
from bs4 import BeautifulSoup
from serpapi import GoogleSearch
import google.generativeai as genai
from docx import Document

from dotenv import load_dotenv
import re


load_dotenv()

SERPAPI_KEY = os.getenv("SERPAPI_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel("gemini-2.5-flash")


def get_ai_overview_urls(keyword):

    params = {
        "engine": "google",
        "q": keyword,
        "api_key": SERPAPI_KEY,
        "gl": "in",    
        "hl": "en"
    }

    search = GoogleSearch(params)
    results = search.get_dict()

    print("DEBUG RESULT:", results)  

    ai_overview = results.get("ai_overview")
    urls = []
    if ai_overview:
        sources = ai_overview.get("sources", [])
        urls = [s["link"] for s in sources if "link" in s]

    if not urls:
        organic = results.get("organic_results", [])
        urls = [o["link"] for o in organic[:5] if "link" in o]

    return urls

def scrape_content(url):

    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        r = requests.get(url, headers=headers)
        soup = BeautifulSoup(r.text, "html.parser")
        paragraphs = [p.get_text() for p in soup.find_all("p")]
        text = " ".join(paragraphs)
        word_count = len(text.split())

        h1 = len(soup.find_all("h1"))
        h2 = len(soup.find_all("h2"))
        h3 = len(soup.find_all("h3"))

        bullet_lists = len(soup.find_all("ul"))
        numbered_lists = len(soup.find_all("ol"))
        tables = len(soup.find_all("table"))

        faq = False
        faq_keywords = ["faq", "faqs", "frequently asked questions"]
        for tag in soup.find_all(["h1", "h2", "h3", "h4"]):
            heading = tag.get_text().lower().strip()
            if any(k in heading for k in faq_keywords):
                faq = True
                break

        if not faq:
            for script in soup.find_all("script", type="application/ld+json"):
                if "FAQPage" in script.text:
                    faq = True
                    break
        comparison_section = False
        if "compare" in text.lower() or "comparison" in text.lower():
            comparison_section = True

        calculators = False
        if soup.find("input") and ("calculate" in text.lower() or "calculator" in text.lower()):
            calculators = True

        return {
            "url": url,
            "text": text[:4000], 
            "word_count": word_count,
            "h1": h1,
            "h2": h2,
            "h3": h3,
            "bullet_lists": bullet_lists,
            "numbered_lists": numbered_lists,
            "tables": tables,
            "faq": faq,
            "comparison_section": comparison_section,
            "calculators": calculators
        }

    except Exception as e:
        print("Error scraping:", url)
        return None

def analyze_gap(ai_pages, client_page):

    prompt = f"""
You are an SEO analyst.
Compare the AI Overview source articles with the client article.
AI Overview Pages:
{ai_pages}
Client Article:
{client_page}
Identify:
1. Missing sections
2. Missing content formats (FAQ, lists, tables)
3. Structural differences
4. Content depth gaps
Then provide:
Gap Summary
Recommendations (3–5 actionable suggestions)
"""

    response = model.generate_content(prompt)

    return response.text

def create_report(keyword, ai_urls, ai_pages, client_page, analysis):
    doc = Document()
    doc.add_heading("AI Overview Content Gap Report", level=1)
    doc.add_heading("Keyword", level=2)
    doc.add_paragraph(keyword)
    doc.add_heading("AI Overview Sources", level=2)
    for url in ai_urls:
        doc.add_paragraph(url)
    doc.add_heading("AI Overview Content Analysis", level=2)
    for page in ai_pages:
        if page:
            doc.add_paragraph(
                f"""
URL: {page['url']}
Word Count: {page['word_count']}
H1: {page['h1']} | H2: {page['h2']} | H3: {page['h3']}
Bullet Lists: {page['bullet_lists']} | Numbered Lists: {page['numbered_lists']}
Tables: {page['tables']} | FAQ detected: {page['faq']}
Comparison Section: {page['comparison_section']} | Calculators: {page['calculators']}
"""
            )

    doc.add_heading("Client Article Analysis", level=2)

    doc.add_paragraph(
        f"""
URL: {client_page['url']}
Word Count: {client_page['word_count']}
H1: {client_page['h1']} | H2: {client_page['h2']} | H3: {client_page['h3']}
Bullet Lists: {client_page['bullet_lists']} | Numbered Lists: {client_page['numbered_lists']}
Tables: {client_page['tables']} | FAQ detected: {client_page['faq']}
Comparison Section: {client_page['comparison_section']} | Calculators: {client_page['calculators']}
"""
    )

    doc.add_heading("Gap Analysis", level=2)
    clean_analysis = re.sub(r"\*\*(.*?)\*\*", r"\1", analysis) 
    clean_analysis = re.sub(r"\*(.*?)\*", r"\1", clean_analysis) 
    
    for line in clean_analysis.split("\n"):
        line = line.strip()
        if line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line:
            doc.add_paragraph(line)

    doc.save("gap_report.docx")

def main():

    keyword = input("Enter keyword: ")
    client_url = input("Enter client URL: ")

    print("\nFetching AI Overview sources...")
    ai_urls = get_ai_overview_urls(keyword)
    if not ai_urls:
        print("No AI Overview sources found. Try another keyword.")
        return

    print("\nScraping AI Overview pages...")
    ai_pages = []

    for url in ai_urls:
        page = scrape_content(url)
        if page:
            ai_pages.append(page)

    print("\nScraping client page.")
    client_page = scrape_content(client_url)

    print("\nRunning AI analysis.")
    analysis = analyze_gap(ai_pages, client_page)

    print("\nGenerating report")
    create_report(keyword, ai_urls, ai_pages, client_page, analysis)

    print("\nreport saved as gap_report.docx")


if __name__ == "__main__":
    main()